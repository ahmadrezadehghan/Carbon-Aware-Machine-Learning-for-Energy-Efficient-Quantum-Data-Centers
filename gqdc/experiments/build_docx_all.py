
# gqdc/experiments/build_docx_all.py
"""
Build a comprehensive Persian (RTL) Word report including all new figures/stats.
Outputs: outputs/gqdc_results_fullpaper.docx
Requires: python-docx
"""
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path
import json, pandas as pd

DEFAULT_FONT = "B Nazanin"

def add_p(doc, txt, size=12, bold=False, rtl=True):
    p = doc.add_paragraph()
    r = p.add_run(txt); r.bold=bold; r.font.size=Pt(size)
    try: r.font.name = DEFAULT_FONT
    except: pass
    if rtl:
        p.paragraph_format.right_to_left = True
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    return p

def add_h(doc, txt): return add_p(doc, txt, size=16, bold=True, rtl=True)

def add_img(doc, path, caption=None, w=5.8):
    p = Path(path)
    if p.exists():
        doc.add_picture(str(p), width=Inches(w))
        if caption:
            c = doc.add_paragraph(caption); c.alignment=WD_ALIGN_PARAGRAPH.CENTER

def maybe_num(doc):
    # facility CI
    p = Path('outputs/facility_summary_ci.csv')
    if p.exists():
        df = pd.read_csv(p)
        row = df[df['metric']=='reduction_pct']
        if not row.empty:
            add_h(doc, "📌 خلاصهٔ Facility")
            add_p(doc, f"کاهش انرژی: {row['mean'].values[0]:.2f}٪  (CI95%: {row['ci95_lo'].values[0]:.2f} تا {row['ci95_hi'].values[0]:.2f})")
    # bootstrap
    j = Path('outputs/bootstrap_energy.json')
    if j.exists():
        d = json.loads(j.read_text(encoding='utf-8'))
        add_h(doc, "📌 Bootstrap")
        add_p(doc, f"میانگین کاهش: {d['mean_reduction_pct']:.2f}٪ | p≈{d['p_value']:.4f} | CI95%: {d['ci95_lo']:.2f}–{d['ci95_hi']:.2f}")
    # stress
    s = Path('outputs/stress_summary.csv')
    if s.exists():
        add_h(doc, "📌 سناریوهای استرس")
        add_p(doc, pd.read_csv(s).to_string(index=False), size=10)

def main():
    doc = Document()
    add_h(doc, "مرکز دادهٔ کوانتومی سبز 🌿⚛️ — گزارش کامل")
    add_p(doc, "این سند مجموعهٔ کامل شکل‌ها و آمارهای مقاله را یک‌جا جمع می‌کند.", 12)

    maybe_num(doc)

    add_h(doc, "شکل‌ها")
    add_img(doc, "outputs/fig_facility_energy_bar.png", "انرژی facility — baseline")
    add_img(doc, "outputs/fig_pareto_energy_wait.png", "پارتو: انرژی در برابر انتظار")
    add_img(doc, "outputs/fig_pareto_energy_sla.png", "پارتو: انرژی در برابر SLA Miss")
    add_img(doc, "outputs/fig_cop_ablation.png", "COP خطی در برابر غیرخطی")
    add_img(doc, "outputs/fig_fairness_violin.png", "عدالت: ویولین انتظار")
    add_img(doc, "outputs/fig_scheduler_bar.png", "مقایسه زمان‌بند: FIFO vs EDF")
    add_img(doc, "outputs/fig_cost_bar.png", "مقایسه هزینه: Fixed vs MPC")
    add_img(doc, "outputs/fig_stress_bar.png", "کاهش انرژی در سناریوهای استرس")
    add_img(doc, "outputs/fig_bootstrap_hist.png", "Bootstrap توزیع میانگین کاهش")

    Path('outputs').mkdir(exist_ok=True)
    out = Path('outputs/gqdc_results_fullpaper.docx')
    doc.save(str(out))
    print("Saved:", out)

if __name__ == '__main__':
    main()
