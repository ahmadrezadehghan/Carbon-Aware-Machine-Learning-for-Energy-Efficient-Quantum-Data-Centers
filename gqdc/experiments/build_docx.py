
# gqdc/experiments/build_docx.py  (UPDATED to include new figures if present)
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path
import pandas as pd

DEFAULT_FONT = "B Nazanin"

def add_p(doc, text, size=12, bold=False, rtl=True):
    p = doc.add_paragraph()
    r = p.add_run(text); r.bold=bold; r.font.size=Pt(size)
    try:
        r.font.name = DEFAULT_FONT
    except Exception: pass
    if rtl:
        p.paragraph_format.right_to_left = True
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    return p

def add_h(doc, text): return add_p(doc, text, size=16, bold=True, rtl=True)

def add_img(doc, path, caption=None, w=5.8):
    p = Path(path)
    if p.exists():
        doc.add_picture(str(p), width=Inches(w))
        if caption:
            c = doc.add_paragraph(caption); c.alignment=WD_ALIGN_PARAGRAPH.CENTER

def maybe_metrics(doc):
    p = Path("outputs/facility_summary_ci.csv")
    if p.exists():
        df = pd.read_csv(p)
        row = df[df['metric']=='reduction_pct']
        if not row.empty:
            red = row['mean'].values[0]; lo=row['ci95_lo'].values[0]; hi=row['ci95_hi'].values[0]
            add_h(doc, "📌 خلاصهٔ عددی (Facility)")
            add_p(doc, f"کاهش انرژی Facility: {red:.2f}٪  (CI95%: {lo:.2f} تا {hi:.2f})")
    p2 = Path("outputs/pareto_points.csv")
    if p2.exists():
        add_h(doc, "📌 خلاصهٔ پارتو")
        add_p(doc, "مرز سازش میان صرفه‌جویی انرژی و انتظار/SLA از sweep شدت تعویق به‌دست آمده است.")

def main():
    doc = Document()
    add_h(doc, "مرکز دادهٔ کوانتومی سبز 🌿⚛️ — گزارش تکمیلی")
    add_p(doc, "این نسخه شامل نتایج پارتو، ابلیشن COP غیرخطی و عدالت بین طبقات است.", 12)

    maybe_metrics(doc)

    add_h(doc, "شکل‌ها (جدید)")
    add_img(doc, "outputs/fig_pareto_energy_wait.png", "پارتو: صرفه‌جویی انرژی در برابر انتظار")
    add_img(doc, "outputs/fig_pareto_energy_sla.png", "پارتو: صرفه‌جویی انرژی در برابر SLA Miss")
    add_img(doc, "outputs/fig_cop_ablation.png", "ابلیشن: COP خطی در برابر غیرخطی")
    add_img(doc, "outputs/fig_fairness_violin.png", "عدالت: توزیع انتظار بین کلاس‌ها")

    Path('outputs').mkdir(exist_ok=True)
    out = Path('outputs/gqdc_results_report_extra.docx')
    doc.save(str(out))
    print("Saved:", out)

if __name__ == '__main__':
    main()
